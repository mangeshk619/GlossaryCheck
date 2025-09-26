import streamlit as st
import pandas as pd
import re, io, os
from collections import Counter, defaultdict
from typing import Optional, List, Tuple

# ---- OpenCC (for Simplified <-> Traditional) ----
try:
    from opencc import OpenCC
    OPENCC = True
    cc_s2t = OpenCC('s2t')  # Simplified -> Traditional
except Exception:
    OPENCC = False
    cc_s2t = None

st.set_page_config(page_title="TW MoJ — Glossary Adherence & Simplified Audit", layout="wide")
st.title("🇹🇼 TW MoJ — Corpus-level Glossary Adherence (EN ↔ ZH-TW) + Simplified Audit")

with st.expander("How it works"):
    st.markdown("""
- Separate corpora; **no alignment** needed.
- **Direction**:
  - **EN → ZH-TW**: enforce Chinese terms in Target; optional filter by English Source.
  - **ZH-TW → EN**: enforce English terms in Target; optional filter by Chinese Source.
- **Glossary**: upload CSV/XLSX or auto-load `glossary_template.csv` from the repo. You explicitly map which column is **English** and which is **Chinese**.
- **Report schema**:
  - EN → ZH-TW: `en_term, zh_tw_expected, adhered_in_target, match_count, matched_variant_examples`
  - ZH-TW → EN: `zh_tw_term, en_expected, adhered_in_target, match_count, matched_variant_examples`
- Optional: `en_term_regex`, `zh_tw_term_regex`, `notes`, multi-variant with `A|B|C`.
- Extras: Simplified-character audit (Chinese corpus) + Mainland→Taiwan term scan.
    """)

# ---- Mainland -> Taiwan seed map (override via CSV if needed) ----
MAINLAND_TO_TW = {
    "软件": "軟體","硬件": "硬體","互联网": "網際網路","网络": "網路","手机": "手機",
    "邮箱": "電子郵件","邮件": "郵件","图标": "圖示","应用程序": "應用程式","服务器": "伺服器",
    "高清": "高畫質","视频": "影片","打印机": "印表機","鼠标": "滑鼠","键盘": "鍵盤",
    "用户": "使用者","复印": "影印","登录": "登入","登出": "登出",
}

# ---------------- Utilities ----------------
_FW_TO_HW = str.maketrans({
    "，": ",","。": ".","；": ";","：": ":","！": "!","？": "?",
    "（": "(","）": ")","【": "[","】": "]","「": '"',"」": '"',
    "『": '"',"』": '"',"、": ",","　": " ","－": "-","～": "~","《": "<","》": ">"
})
def normalize_zh(s: str) -> str:
    if s is None: return ""
    s = str(s).translate(_FW_TO_HW)
    s = re.sub(r"\s+", " ", s).strip()
    return s

def normalize_en(s: str) -> str:
    if s is None: return ""
    return re.sub(r"\s+", " ", str(s)).strip()

def _read_docx(file) -> str:
    from docx import Document
    d = Document(file)
    parts = []
    for p in d.paragraphs:
        parts.append(p.text)
    for table in d.tables:
        for row in table.rows:
            for cell in row.cells:
                parts.append(cell.text)
    return "\n".join([p for p in parts if p is not None])

def read_any(file) -> Tuple[str, List[str]]:
    name = file.name.lower()
    if name.endswith(".txt"):
        raw = file.read().decode("utf-8", errors="ignore")
        segs = [ln.strip() for ln in raw.splitlines()]
        return raw, segs
    elif name.endswith(".csv"):
        df = pd.read_csv(file, dtype=str).fillna("")
        segs = [" ".join(map(str, r)) for r in df.values]
        return "\n".join(segs), segs
    elif name.endswith(".tsv"):
        df = pd.read_csv(file, sep="\t", dtype=str).fillna("")
        segs = [" ".join(map(str, r)) for r in df.values]
        return "\n".join(segs), segs
    elif name.endswith((".xlsx",".xls")):
        df = pd.read_excel(file, dtype=str).fillna("")
        segs = [" ".join(map(str, r)) for r in df.values]
        return "\n".join(segs), segs
    elif name.endswith(".docx"):
        raw = _read_docx(file)
        segs = [ln.strip() for ln in raw.splitlines()]
        return raw, segs
    else:
        st.error(f"Unsupported file type: {name}")
        return "", []

def load_table(file) -> Optional[pd.DataFrame]:
    name = file.name.lower()
    if name.endswith(".csv"):
        return pd.read_csv(file, dtype=str).fillna("")
    if name.endswith(".tsv"):
        return pd.read_csv(file, sep="\t", dtype=str).fillna("")
    if name.endswith((".xlsx",".xls")):
        return pd.read_excel(file, dtype=str).fillna("")
    return None

# ---------- Glossary helpers ----------
def read_glossary_raw(gls_file) -> pd.DataFrame:
    if gls_file.name.lower().endswith(".csv"):
        return pd.read_csv(gls_file, dtype=str).fillna("")
    else:
        return pd.read_excel(gls_file, dtype=str).fillna("")

def read_repo_glossary_template() -> Optional[pd.DataFrame]:
    path = os.path.join(os.getcwd(), "glossary_template.csv")
    if os.path.exists(path):
        try:
            df = pd.read_csv(path, dtype=str).fillna("")
            df._source_path = path  # optional
            return df
        except Exception:
            return None
    return None

def load_glossary_with_user_mapping_df(df: pd.DataFrame, user_en_col: str, user_zh_col: str) -> pd.DataFrame:
    cols = list(df.columns)
    if user_en_col not in cols or user_zh_col not in cols:
        raise ValueError("Please choose valid columns for English and Chinese terms.")
    out = df.rename(columns={user_en_col: "en_term", user_zh_col: "zh_tw_term"})
    # carry optional columns if present
    for c in df.columns:
        cl = c.strip().lower()
        if cl == "en_term_regex" and "en_term_regex" not in out.columns:
            out["en_term_regex"] = df[c]
        if cl == "zh_tw_term_regex" and "zh_tw_term_regex" not in out.columns:
            out["zh_tw_term_regex"] = df[c]
        if cl == "notes" and "notes" not in out.columns:
            out["notes"] = df[c]
    return out

def split_variants(v: str) -> List[str]:
    if not v: return []
    return [x.strip() for x in str(v).split("|") if x.strip()]

def sample_contexts(text: str, needle: str, max_hits=5, window=32):
    ctxs, start = [], 0
    while len(ctxs) < max_hits:
        pos = text.find(needle, start)
        if pos == -1: break
        ctxs.append(text[max(0,pos-window):pos+len(needle)+window].replace("\n"," "))
        start = pos + len(needle)
    return ctxs

# ---------- EN target literal matcher (case-insensitive, optional whole-word) ----------
def en_literal_find_all(hay: str, needle: str, whole_word: bool) -> Tuple[int, List[str]]:
    if not needle.strip():
        return 0, []
    pat = re.escape(needle.strip())
    if whole_word:
        pat = r"\b" + pat + r"\b"
    regex = re.compile(pat, re.IGNORECASE)
    matches = list(regex.finditer(hay))
    contexts = []
    for m in matches[:5]:
        a, b = m.start(), m.end()
        contexts.append(hay[max(0,a-32):b+32].replace("\n"," "))
    return len(matches), contexts

# ---- Simplified Character Audit ----
def simplified_char_audit(raw_text: str, segments: List[str], context_window: int = 30,
                          ignore_chars: Optional[set] = None):
    if not OPENCC:
        return pd.DataFrame(), pd.DataFrame()

    ignore_chars = ignore_chars or set()
    joined = "\n".join(segments)
    seg_starts, pos = [], 0
    for s in segments:
        seg_starts.append(pos)
        pos += len(s) + 1

    def find_seg(abs_pos: int) -> int:
        lo, hi = 0, len(seg_starts) - 1
        while lo <= hi:
            mid = (lo + hi) // 2
            if seg_starts[mid] <= abs_pos:
                lo = mid + 1
            else:
                hi = mid - 1
        return max(0, hi)

    sim_occurrences = []
    for idx, ch in enumerate(joined):
        if re.match(r"[\u4e00-\u9fff]", ch):
            if cc_s2t.convert(ch) != ch and ch not in ignore_chars:
                sim_occurrences.append((ch, idx))

    char_counts = Counter([c for c, _ in sim_occurrences])
    char_contexts = defaultdict(list)
    occ_rows = []
    for ch, abs_pos in sim_occurrences:
        seg_idx = find_seg(abs_pos)
        left = max(0, abs_pos - context_window)
        right = min(len(joined), abs_pos + context_window)
        ctx = joined[left:right].replace("\n", " ")
        sugg = cc_s2t.convert(ch)
        occ_rows.append({
            "segment_index": seg_idx + 1,
            "character": ch,
            "suggested_traditional": sugg,
            "context": ctx
        })
        if len(char_contexts[ch]) < 5:
            char_contexts[ch].append(ctx)

    sum_rows = []
    for ch, cnt in sorted(char_counts.items(), key=lambda x: (-x[1], x[0])):
        sum_rows.append({
            "character": ch,
            "count": cnt,
            "suggested_traditional": cc_s2t.convert(ch),
            "sample_contexts": " … ".join(char_contexts[ch])
        })

    return pd.DataFrame(sum_rows), pd.DataFrame(occ_rows)

# ---------------- Sidebar ----------------
st.sidebar.header("Uploads")
src_file = st.sidebar.file_uploader("Optional Source corpus", type=["txt","csv","tsv","xlsx","xls","docx"])
tgt_file = st.sidebar.file_uploader("Target corpus (required for run)", type=["txt","csv","tsv","xlsx","xls","docx"])
gls_file = st.sidebar.file_uploader("Glossary (CSV/XLSX)", type=["csv","xlsx","xls"])

st.sidebar.header("Direction & Mode")
direction = st.sidebar.radio("Direction", ["EN → ZH-TW","ZH-TW → EN"], index=0)
mode = st.sidebar.radio("Check mode", ["Filtered by Source (recommended)","Target-only"], index=0)

st.sidebar.header("Matching (EN side)")
whole_word = st.sidebar.checkbox("Whole-word when matching EN", True)  # used for EN filtering and EN target matching
case_sensitive = st.sidebar.checkbox("EN filter is case-sensitive", False)

st.sidebar.header("Simplified Audit Options")
context_window = st.sidebar.slider("Context window (chars around hit)", 10, 80, 30, step=2)
ignore_input = st.sidebar.text_input("Ignore these Chinese characters (comma-separated)", value="")
ignore_chars = set([c.strip() for c in ignore_input.split(",") if c.strip()])

# ---- Load overrides for Mainland→Taiwan map ----
override_file = st.sidebar.file_uploader("Optional Mainland→Taiwan overrides (CSV/XLSX)", type=["csv","xlsx","xls"])
if override_file is not None:
    try:
        df_o = load_table(override_file)
        add_map = {str(r["mainland"]).strip(): str(r["taiwan"]).strip()
                   for _, r in df_o.iterrows() if str(r.get("mainland","")).strip()}
        MAINLAND_TO_TW.update(add_map)
        st.sidebar.success(f"Loaded {len(add_map)} Mainland→Taiwan overrides.")
    except Exception as e:
        st.sidebar.warning(f"Override load failed: {e}")

# ---- Glossary: upload or auto-load from repo ----
st.sidebar.header("Glossary source")
use_repo_template = st.sidebar.checkbox("Use glossary_template.csv from repo if no upload", value=True)

gls_df_preview = None
glossary_source = None
if gls_file is not None:
    gls_df_preview = read_glossary_raw(gls_file)
    glossary_source = "uploaded"
elif use_repo_template:
    repo_df = read_repo_glossary_template()
    if repo_df is not None:
        gls_df_preview = repo_df
        glossary_source = "repo_template"

# If we have any glossary df, let user map columns
en_col_choice = zh_col_choice = None
if gls_df_preview is not None:
    st.sidebar.subheader("Glossary column mapping")
    cols = list(gls_df_preview.columns)

    def guess_index(names, fallbacks):
        for fb in fallbacks:
            for i, c in enumerate(names):
                if c.strip().lower() == fb:
                    return i
        return 0

    if glossary_source == "repo_template":
        en_default = guess_index(cols, ["en_term","english","en"])
        zh_default = guess_index(cols, ["zh_tw_term","zh_tw","zh-hant","traditional_chinese","zh"])
    else:
        en_default = guess_index(cols, ["en_term","english","en","target_en","expected_en","eng"])
        zh_default = guess_index(cols, ["zh_tw_term","zh_tw","zh-hant","traditional_chinese","zh","chinese","source_zh"])

    en_col_choice = st.sidebar.selectbox("English term column", cols, index=en_default)
    zh_col_choice = st.sidebar.selectbox("Chinese term column (Traditional)", cols, index=zh_default)

# ---------------- RUN ----------------
if tgt_file and (gls_df_preview is not None) and (en_col_choice and zh_col_choice):
    # Read corpora
    def read_any_public(file):
        if not file: return ("", [])
        name = file.name.lower()
        if name.endswith(".txt"):
            raw = file.read().decode("utf-8", errors="ignore")
            segs = [ln.strip() for ln in raw.splitlines()]
            return raw, segs
        elif name.endswith(".csv"):
            df = pd.read_csv(file, dtype=str).fillna("")
            segs = [" ".join(map(str, r)) for r in df.values]
            return "\n".join(segs), segs
        elif name.endswith(".tsv"):
            df = pd.read_csv(file, sep="\t", dtype=str).fillna("")
            segs = [" ".join(map(str, r)) for r in df.values]
            return "\n".join(segs), segs
        elif name.endswith((".xlsx",".xls")):
            df = pd.read_excel(file, dtype=str).fillna("")
            segs = [" ".join(map(str, r)) for r in df.values]
            return "\n".join(segs), segs
        elif name.endswith(".docx"):
            raw = _read_docx(file)
            segs = [ln.strip() for ln in raw.splitlines()]
            return raw, segs
        else:
            return "", []

    tgt_raw, tgt_segments = read_any_public(tgt_file)
    src_raw, src_segments = read_any_public(src_file)

    # Normalize per direction
    if direction == "EN → ZH-TW":
        src_norm = normalize_en(src_raw) if src_raw else ""
        tgt_norm = normalize_zh(tgt_raw)
        chinese_raw = tgt_raw
        chinese_segments = tgt_segments
    else:  # ZH-TW → EN
        src_norm = normalize_zh(src_raw) if src_raw else ""
        tgt_norm = normalize_en(tgt_raw)
        chinese_raw = src_raw if src_file else ""
        chinese_segments = src_segments if src_file else []

    # Load glossary with explicit mapping
    try:
        glossary = load_glossary_with_user_mapping_df(gls_df_preview, en_col_choice, zh_col_choice)
    except Exception as e:
        st.error(str(e)); st.stop()

    # Choose which columns filter vs expected (by direction)
    if direction == "EN → ZH-TW":
        filter_term_col, filter_rx_col = "en_term", "en_term_regex"          # EN filter (optional)
        expected_term_col, expected_rx_col = "zh_tw_term", "zh_tw_term_regex"  # ZH expected in TARGET
    else:
        filter_term_col, filter_rx_col = "zh_tw_term", "zh_tw_term_regex"    # ZH filter (optional)
        expected_term_col, expected_rx_col = "en_term", "en_term_regex"      # EN expected in TARGET

    # Filter-by-Source (optional)
    if mode == "Filtered by Source (recommended)":
        if not src_file:
            st.warning("Filtered-by-Source selected, but no Source uploaded. All glossary rows will be enforced.")
        keep = []
        for _, r in glossary.iterrows():
            term = str(r.get(filter_term_col, ""))
            is_rx = str(r.get(filter_rx_col, "")).strip().lower() in ("y","yes","true","1")
            if not term:
                keep.append(False); continue
            if not src_norm:
                keep.append(True); continue
            if direction == "EN → ZH-TW":
                flags = 0 if case_sensitive else re.IGNORECASE
                if is_rx:
                    try: pat = re.compile(term, flags)
                    except re.error: keep.append(False); continue
                    keep.append(bool(pat.search(src_norm)))
                else:
                    pat = re.escape(term)
                    if whole_word: pat = r"\b"+pat+r"\b"
                    keep.append(bool(re.search(pat, src_norm, flags)))
            else:
                if is_rx:
                    try: pat = re.compile(term)
                    except re.error: keep.append(False); continue
                    keep.append(bool(pat.search(src_norm)))
                else:
                    keep.append(normalize_zh(term) in src_norm)
        glossary = glossary.loc[keep].reset_index(drop=True)
        st.info(f"Enforcing {len(glossary)} glossary rows (filtered by Source corpus).")
    else:
        st.info(f"Target-only mode: enforcing all {len(glossary)} rows on Target corpus.")

    # ========= Enforce on Target (collect rows; keep BOTH sides) =========
    rows = []
    for _, r in glossary.iterrows():
        zh_gloss = str(r.get("zh_tw_term", ""))   # always kept
        en_gloss = str(r.get("en_term", ""))      # always kept

        expected = str(r.get(expected_term_col, ""))
        is_rx = str(r.get(expected_rx_col, "")).strip().lower() in ("y","yes","true","1")

        match_count, matched_variants = 0, []

        if is_rx:
            # Regex: case-insensitive if EN is target; default otherwise
            flags = re.IGNORECASE if direction == "ZH-TW → EN" else 0
            try:
                pat = re.compile(expected, flags)
            except re.error:
                rows.append({
                    "zh_tw_term": zh_gloss,
                    "en_term": en_gloss,
                    "adhered_in_target": False,
                    "match_count": 0,
                    "matched_variant_examples": "",
                })
                continue
            hits = list(pat.finditer(tgt_norm))
            match_count = len(hits)
            if match_count > 0:
                matched_variants.append(expected)
        else:
            # Literal; support variants A|B|C on the expected side
            variants = split_variants(expected) if expected else []
            if not variants:
                variants = [expected]

            for v in variants:
                if direction == "EN → ZH-TW":
                    # Chinese literal: normalize & contains
                    vv = normalize_zh(v)
                    if not vv: continue
                    start = 0
                    while True:
                        pos = tgt_norm.find(vv, start)
                        if pos == -1: break
                        match_count += 1
                        matched_variants.append(v)
                        start = pos + len(vv)
                else:
                    # English literal: case-insensitive + optional whole-word
                    pat = re.escape(v.strip())
                    if whole_word:
                        pat = r"\b" + pat + r"\b"
                    hits = list(re.finditer(pat, tgt_norm, flags=re.IGNORECASE))
                    if hits:
                        match_count += len(hits)
                        matched_variants.append(v)

        rows.append({
            "zh_tw_term": zh_gloss,
            "en_term": en_gloss,
            "adhered_in_target": match_count > 0,
            "match_count": match_count,
            "matched_variant_examples": "|".join(sorted(set(matched_variants))) if matched_variants else "",
        })

    # ========= Direction-specific report schema =========
    if direction == "EN → ZH-TW":
        report_df = pd.DataFrame([{
            "en_term": r["en_term"],
            "zh_tw_expected": r["zh_tw_term"],
            "adhered_in_target": r["adhered_in_target"],
            "match_count": r["match_count"],
            "matched_variant_examples": r["matched_variant_examples"],
        } for r in rows], columns=["en_term","zh_tw_expected","adhered_in_target","match_count","matched_variant_examples"])
    else:
        report_df = pd.DataFrame([{
            "zh_tw_term": r["zh_tw_term"],
            "en_expected": r["en_term"],
            "adhered_in_target": r["adhered_in_target"],
            "match_count": r["match_count"],
            "matched_variant_examples": r["matched_variant_examples"],
        } for r in rows], columns=["zh_tw_term","en_expected","adhered_in_target","match_count","matched_variant_examples"])

    # ========= Simplified audit & Mainland terms on the Chinese corpus =========
    if direction == "EN → ZH-TW":
        chinese_corpus = tgt_raw
        chinese_segments = tgt_segments
    else:
        chinese_corpus = src_raw if src_file else ""
        chinese_segments = src_segments if src_file else []

    if chinese_corpus:
        if OPENCC:
            simp_summary_df, simp_occ_df = simplified_char_audit(
                chinese_corpus, chinese_segments, context_window=context_window, ignore_chars=ignore_chars
            )
        else:
            simp_summary_df, simp_occ_df = pd.DataFrame(), pd.DataFrame()
        ml_hits = []
        for ml, tw in MAINLAND_TO_TW.items():
            if ml in chinese_corpus:
                ml_hits.append({
                    "mainland_term": ml,
                    "suggested_tw": tw,
                    "sample_contexts": " … ".join(sample_contexts(chinese_corpus, ml))
                })
        ml_df = pd.DataFrame(ml_hits) if ml_hits else pd.DataFrame([{"info":"No mainland terms (from current list) detected"}])
    else:
        simp_summary_df, simp_occ_df = pd.DataFrame(), pd.DataFrame()
        ml_df = pd.DataFrame([{"info":"No Chinese corpus to check"}])

    # ========= KPIs & tables =========
    c1, c2, c3 = st.columns(3)
    enforced = len(report_df)
    adhered = int(report_df["adhered_in_target"].sum()) if enforced else 0
    c1.metric("Glossary rows enforced", enforced)
    c2.metric("Adhered in Target", adhered)
    c3.metric("Not adhered", enforced - adhered)

    st.subheader("Glossary adherence (required schema)")
    st.dataframe(report_df, use_container_width=True)

    st.subheader("Simplified Character Audit")
    if not OPENCC:
        st.warning("OpenCC not available. Add `opencc-python-reimplemented` to requirements.txt to enable Simplified audit.")
    colA, colB = st.columns(2)
    with colA:
        st.markdown("**Summary by character**")
        st.dataframe(simp_summary_df if not simp_summary_df.empty else pd.DataFrame([{"info":"No simplified chars flagged"}]), use_container_width=True)
    with colB:
        st.markdown("**All occurrences (segment-level)**")
        st.dataframe(simp_occ_df if not simp_occ_df.empty else pd.DataFrame([{"info":"No occurrences"}]), use_container_width=True)

    st.subheader("Mainland terms in Chinese corpus")
    st.dataframe(ml_df, use_container_width=True)

    # Basic stats (Target)
    def text_stats(s: str, segments: List[str]):
        cjk = re.findall(r'[\u4e00-\u9fff]', s)
        latin = re.findall(r'[A-Za-z]', s)
        digits = re.findall(r'[0-9]', s)
        return pd.DataFrame([{
            "chars_total": len(s),
            "cjk_chars": len(cjk),
            "latin_chars": len(latin),
            "digit_chars": len(digits),
            "segments": len(segments),
        }])
    stats_df = text_stats(tgt_raw, tgt_segments)

    # ========= Export to Excel (first sheet = report_df) =========
    out = io.BytesIO()
    with pd.ExcelWriter(out, engine="xlsxwriter") as xw:
        # Name the first sheet consistently
        sheet_name = "coverage_minimal"
        report_df.to_excel(xw, index=False, sheet_name=sheet_name)
        (simp_summary_df if not simp_summary_df.empty else pd.DataFrame([{"info":"No simplified chars flagged"}])).to_excel(xw, index=False, sheet_name="simplified_summary")
        (simp_occ_df if not simp_occ_df.empty else pd.DataFrame([{"info":"No occurrences"}])).to_excel(xw, index=False, sheet_name="simplified_occurrences")
        ml_df.to_excel(xw, index=False, sheet_name="mainland_vs_tw")
        stats_df.to_excel(xw, index=False, sheet_name="text_stats")
        pd.DataFrame([{
            "direction": direction, "mode": mode, "whole_word_EN": whole_word,
            "glossary_source": ("uploaded" if st.session_state.get("dummy", False) or ('gls_file' in locals() and gls_file) else (glossary_source or "unknown")),
            "en_col": en_col_choice, "zh_col": zh_col_choice
        }]).to_excel(xw, index=False, sheet_name="run_info")

    st.download_button("Download report.xlsx", out.getvalue(), "twmoj_corpus_report.xlsx",
                       mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet")

else:
    st.info("Upload **Target corpus** and a **Glossary** (or auto-load `glossary_template.csv`), then choose the **English** and **Chinese** columns in the sidebar. (Source is optional for Filtered-by-Source.)")
